import streamlit as st
from google import genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- ページ設定 ---
st.set_page_config(page_title="議事録整形ツール", page_icon="📝")

# Geminiクライアント初期化
try:
    # secrets.toml からAPIキーを取得
    client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
except Exception:
    st.error("APIキーが設定されていません。管理者に相談してください（塩入）。")

def extract_text_from_docx(file):
    """アップロードされたWordファイルからテキストを抽出"""
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def create_formatted_docx(gemini_text):
    """Geminiの回答を解析してWordファイルを生成"""
    new_doc = Document()
    
    # 標準フォント設定（游明朝）
    style = new_doc.styles['Normal']
    style.font.name = '游明朝'
    style.font.size = Pt(10.5)

    lines = gemini_text.split('\n')
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue

        # 特殊タグの判定と書き込み
        if clean_line.startswith('[TITLE]'):
            p = new_doc.add_heading(clean_line.replace('[TITLE]', '').strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif clean_line.startswith('[DATE]'):
            p = new_doc.add_paragraph(clean_line.replace('[DATE]', '').strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        elif clean_line.startswith('[MEMBERS]'):
            member_text = clean_line.replace('[MEMBERS]', '').strip()
            p = new_doc.add_paragraph(f"出席者：{member_text}")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            new_doc.add_paragraph() # 議題の前に1行空ける
            
        elif clean_line.startswith('[H1]'):
            new_doc.add_heading(clean_line.replace('[H1]', '').strip(), level=1)
            
        elif clean_line.startswith('L1:'):
            text = clean_line.replace('L1:', '').strip()
            new_doc.add_paragraph(text, style='List Bullet')
            
        elif clean_line.startswith('L2:'):
            text = clean_line.replace('L2:', '').strip()
            p = new_doc.add_paragraph(text, style='List Bullet 2')
            p.paragraph_format.left_indent = Inches(0.4)
            
        elif clean_line.startswith('L3:'):
            text = clean_line.replace('L3:', '').strip()
            p = new_doc.add_paragraph(text, style='List Bullet 3')
            p.paragraph_format.left_indent = Inches(0.8)
            
        else:
            # タグがない場合は通常の段落として追加
            new_doc.add_paragraph(clean_line)
    
    # メモリ上に保存
    doc_io = io.BytesIO()
    new_doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- UIレイアウト ---
st.title("議事録自動整形ツール（テスト）")
st.write("Wordファイルをドラッグ＆ドロップするだけで、構造化された議事録に整形します。")

# 1. ファイルアップロード
uploaded_file = st.file_uploader(
    "Wordファイル (.docx) を選択、またはここにドラッグしてください", 
    type=["docx"]
)

if uploaded_file:
    # 2. 整形実行ボタン
    if st.button("議事録を整形する", type="primary"):
        with st.spinner("分析して整形中..."):
            try:
                # テキスト抽出
                raw_text = extract_text_from_docx(uploaded_file)
                
                # Geminiへの指示（プロンプト）
                prompt = f"""
                以下の議事録データをビジネス文書として整理してください。
                冒頭に必ず以下の形式で基本情報を入れてください。

                [TITLE]会議名
                [DATE]開催日
                [MEMBERS]出席者

                その後に、以下のルールで内容を続けてください。
                ・議題（見出し）： [H1]議題名
                ・箇条書き1階層： L1:内容
                ・箇条書き2階層：   L2:内容
                ・箇条書き3階層：     L3:内容
                ・**などの記号（太字記号）は絶対に使用しないでください。
                ・「決定事項」や「次回アクション」も階層を使って明確に整理してください。

                【議事録データ】
                {raw_text}
                """

                # Gemini API呼び出し
                response = client.models.generate_content(
                    model="gemini-2.5-flash", 
                    contents=prompt
                )
                
                # Wordファイルを生成
                formatted_doc_io = create_formatted_docx(response.text)
                
                st.success("整形が完了しました！")

                # 3. ダウンロードボタン
                st.download_button(
                    label="📥 整形済みWordファイルをダウンロード",
                    data=formatted_doc_io,
                    file_name=f"整形済み_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                # プレビュー
                with st.expander("出力のテキストを確認"):
                    st.text(response.text)
                    
            except Exception as e:
                st.error(f"処理中にエラーが発生しました: {e}")